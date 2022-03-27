import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from datetime import date
from docx import Document
from docx.shared import Inches

"""
Standard Test Methods for Flexural Properties of Unreinforced and Reinforced Plastics and
"""


class test_report(object):
    """
    General class for ASTMD test reports
    """
    reference = ''
    report_file = ''

    @staticmethod
    def get_test_data(filenames):
        """
        Method reading the test data for every test using the pandas module
        :param filenames: List of test files
        :return: None
        """
        tests = []
        for filename in filenames:
            df = pd.read_csv(filename, delim_whitespace=True, skiprows=5, engine='python')  # read data from file
            df = df.drop(0)  # remove the units
            tests.append(df.astype(float))
        return tests

    @staticmethod
    def average_stress_strain(stresses, strains):
        """
        Compute the average stress strain data from all the stress strain arrays
        :param stresses: List of stress arrays corresponding to tests
        :param strains: List of strain arrays corresponding to tests
        :return: average_stress, average_strain (arrays)
        """
        max_len_stress = np.min([arr.shape[0] for arr in stresses])
        max_len_strain = np.min([arr.shape[0] for arr in strains])

        if max_len_strain != max_len_stress:
            raise ValueError("Not the same amount of points for strain and stress")
        else:
            i = max_len_strain

        avg_stress = np.mean(np.vstack([arr[:i] for arr in stresses]), axis=0)
        avg_strain = np.mean(np.vstack([arr[:i] for arr in strains]), axis=0)

        return avg_stress, avg_strain

    @staticmethod
    def compute_strength(stresses, all_values=False):
        """
        Method computing the rupture modulus of the material from a list of stress arrays
        :param stresses: list of stress arrays
        :return: mean rupture modulus and standard deviation
        """
        # max stress = strength
        strengths = np.array([arr.max() for arr in stresses])
        if all_values:
            return strengths
        else:
            return strengths.mean(), strengths.std()

    @staticmethod
    def compute_stress_at_break(stresses):
        """
        Method computing the average stress at break for the test samples
        :return: avg, std stress at break
        """
        break_stresses = np.array([arr[-1] for arr in stresses])
        return break_stresses.mean(), break_stresses.std()

    @staticmethod
    def compute_strain_at_break_5p100(strains):
        """
        Method computing if the speciens broke below 5% strain
        :param strains: list of strain arrays
        :return: array containing ['yes', 'no'] strings to write in report
        """
        values = []
        for strain in strains:
            if strain[-1] > 0.05:
                values.append('no')
            else:
                values.append('yes')
        return values

    @staticmethod
    def compute_break_strain(strains, all_values=False):
        """
        Method to compute the strain at break
        :param all_values: Return parameter values for all samples
        :param strains: list of strain arrays
        :return: avg, std strain at break
        """
        break_strains = np.array([arr[-1] for arr in strains])
        if all_values:
            return break_strains
        else:
            return break_strains.mean(), break_strains.std()

    @staticmethod
    def compute_cross_head_rate(tests):
        """Method to compute the cross head rate from the test data"""
        rates = []
        for test in tests:
            rate = np.mean(np.gradient(test['Crosshead'], test['Time']))
            rates.append(rate)
        return np.mean(rates)

    @staticmethod
    def make_stress_strain_plot(filename, all_stress, all_strain):
        """
        Makes the png file corresponding to the stress-strain plot required in the report
        :param all_strain: list of strain arrays
        :param all_stress: list of stress arrays
        :param filename: filename used to save the file
        :return: None
        """
        avg_stress, avg_strain = test_report.average_stress_strain(all_stress, all_strain)

        fig, ax = plt.subplots(figsize=(8, 6))

        for i, stress, strain in zip(range(len(all_stress)), all_stress, all_strain):
            if i == 0:
                ax.plot(strain, stress, color='0.75', label='all samples data')
            else:
                ax.plot(strain, stress, color='0.8')

        ax.plot(avg_strain, avg_stress, color='k', label='average curve')
        ax.set_xlabel('strain mm/mm')
        ax.set_ylabel('stress MPa')
        ax.set_title('Stress-strain curve for all samples data and average')
        ax.legend()
        plt.grid()
        fig.savefig(filename, dpi=300)


class D790(test_report):
    """ Class for computing the results of the ASTMD790 test standard
    ASTM International. (2017). Standard Test Methods for Flexural Properties of Unreinforced
    and Reinforced Plastics and Electrical Insulating Materials."""

    reference = "ASTM International. (2017). Standard Test Methods for Flexural Properties of Unreinforced and " \
                "Reinforced Plastics and Electrical Insulating Materials. "

    def __init__(self, filenames, widths, depths, span, mtr_name="", largespan=False):
        """
        Constructor class for the ASTM D790 test standard
        :param filenames: filenames for the test data in a list
        :param widths: List of the sample widths in milimeters
        :param depths: List of the sample depths (thickness) in milimeters
        :param span: Float, Span used to test the samples
        :param mtr_name: String, Name of the material
        :param largespan: Bool, If the special testcase for large span was used
        This class does all the calculations needed for an ASTM D790 test report and writes the results to file
        """
        # Store the usefull attributes
        self.name = mtr_name
        self.widths = widths  # widths of the samples (list)
        self.depths = depths  # depths of the samples (list)
        self.L = span  # span used to test the samples (float)
        self.largespan = largespan
        # Read the test data from the files
        self.tests = self.get_test_data(filenames)
        self.write_report()

    def compute_stress_strain(self):
        """
        Method computing the stress-strain data from the tests results
        Normal span stress formula is from Eq. 3 in ASTM D790 section 12.2
        Large span stress formula is from Eq. 4 in ASTM D790 section 12.3
        Strain formula is from Eq. 5 in ASTM D790 section 12.4
        :return: stress and strain for all the tests
        """
        largespan = self.largespan
        all_stresses = []
        all_strains = []
        for test, b, d in zip(self.tests, self.widths, self.depths):
            if largespan:
                # Eq. 4 in ASTM D790 section 12.3
                a = (3*test['Load']*self.L)/(2*b*d**2)
                stress = a * (1 + 6*(test['Crosshead']/self.L)**2 - 4*(d/self.L)*(test['Crosshead']/self.L))
                all_stresses.append(np.array(stress))
            else:
                # Eq. 3 in ASTM D790 section 12.2
                stress = (3 * test['Load'] * self.L) / (2*b*d**2)
                all_stresses.append(np.array(stress))

            # Eq. 5 in ASTM D790 section 12.4
            strain = (6 * test['Crosshead'] * d) / (self.L ** 2)
            all_strains.append(np.array(strain))

        return all_stresses, all_strains

    def compute_tgt_modulus(self):
        """
        Compute the tangent modulus from Eq. 6 Section 12.5.1 in standard ASTM D790
        :return:
        """
        # empty list to store modulus values
        modulus = []
        for test, b, d in zip(self.tests, self.widths, self.depths):

            # Get max slope in the load - deflection plot
            load = np.array(test['Load'])
            deflexion = np.array(test['Crosshead'])
            deflexion = deflexion[:np.argmax(load)]
            load = load[:np.argmax(load)]
            slopes = list()
            for e in range(0, load.shape[0] - 25, 5):
                slopes.append((load[e + 25] - load[e]) / (deflexion[e + 25] - deflexion[e]))
            m = np.max(slopes)
            # compute the modulus
            modulus.append((self.L ** 3 * m) / (4 * b * d**3))

        modulus = np.array(modulus)
        return modulus.mean(), modulus.std()

    def write_report(self):
        """
        Method writing the ASTM D790 test report to a .docx (microsoft word) file
        :return: None
        """
        # Compute stress and strain values
        stresses, strains = self.compute_stress_strain()

        # Document header
        document = Document()
        document.add_heading('ASTM D790 Test report', 1)
        document.add_heading(date.today().isoformat(), 3)
        document.add_paragraph(self.reference)

        # Material identification
        document.add_heading('Material Identification', 2)
        document.add_paragraph(f'Material name: {self.name}', style='List Bullet')
        document.add_paragraph('Material type: ', style='List Bullet')
        document.add_paragraph('Material source: ', style='List Bullet')
        document.add_paragraph('Material code number: ', style='List Bullet')
        document.add_paragraph('Composite material ply stacking sequence: ', style='List Bullet')

        # Specimen preparation
        document.add_heading('Specimen preparation', 2)
        document.add_paragraph('How the specimens where prepared: ', style='List Bullet')
        document.add_paragraph('Direction of cutting and loading of the specimens: ', style='List Bullet')
        document.add_paragraph('Conditioning procedure: ', style='List Bullet')
        p = document.add_paragraph()
        p.add_run('Specimens dimensions').bold = True
        table = document.add_table(rows=1, cols=3, style='Table Grid')
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Specimen number'
        header_cells[1].text = 'Width (mm)'
        header_cells[2].text = 'Depth (mm)'
        for id, b, d in zip(range(1, len(self.widths)+1), self.widths, self.depths):
            row_cells = table.add_row().cells
            row_cells[0].text = str(id)
            row_cells[1].text = str(b)
            row_cells[2].text = str(d)

        # Testing procedure
        document.add_heading('Testing procedure', 2)
        p = document.add_paragraph()
        p.add_run('Remove the procedure/type you have not used').bold = True
        document.add_paragraph('Procedure A: Rate of cross head motion was computed using Eq. (1) section 10.1.3',
                               style='List Bullet')
        document.add_paragraph('Procedure B: Rate of cross head motion was determined so that the rate of straining '
                               'of the outer surface of the test specimen is 0.10 mm/mm',
                               style='List Bullet')
        document.add_paragraph('Type I: Deflexion was measured using cross head motion',
                               style='List Bullet')
        document.add_paragraph('Type II: Deflexion was measured using a deflectometer',
                               style='List Bullet')
        document.add_paragraph(f'Support span length: {self.L} mm', style='List Bullet')
        if self.largespan:
            document.add_paragraph('Large span was used: Yes', style='List Bullet')
        else:
            document.add_paragraph('Large span was used: No', style='List Bullet')
        if not np.isclose(self.L/np.mean(self.depths), 16):
            ratio = int(np.round(self.L/np.mean(self.depths)))
            document.add_paragraph(f'Support span-to-depth ratio: {ratio}:1', style='List Bullet')
        document.add_paragraph('Radius of loading nose and supports: 5 mm (change value if not 5 mm)',
                               style='List Bullet')
        cross_head_rate = self.compute_cross_head_rate(self.tests)
        rate = int(np.round(cross_head_rate*60))
        document.add_paragraph(f'Rate of cross head motion: {rate} mm/min', style='List Bullet')

        # Stress-strain plot section
        document.add_heading('Flexural stress-strain curves', 2)
        self.make_stress_strain_plot('stress-strain_temp', all_stress=stresses, all_strain=strains)
        document.add_picture('stress-strain_temp.png', width=Inches(6))
        os.remove('stress-strain_temp.png')

        # Test results section
        document.add_heading('Test results', 2)
        document.add_paragraph('Was a speciment rejected after testing: Yes, No, why ?', style='List Bullet')
        p = document.add_paragraph()
        p.add_run('Average properties test results').bold = True
        table = document.add_table(rows=1, cols=3, style='Table Grid')
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Property'
        header_cells[1].text = 'Mean value'
        header_cells[2].text = 'Standard deviation'
        # Modulus
        row_cells = table.add_row().cells
        row_cells[0].text = 'Tangent bending modulus (MPa)'
        avg_modulus, std_modulus = self.compute_tgt_modulus()
        row_cells[1].text = str(int(np.round(avg_modulus)))
        row_cells[2].text = str(int(np.round(std_modulus)))
        # Strength
        row_cells = table.add_row().cells
        row_cells[0].text = 'Flexural strength (MPa)'
        avg_strength, std_strength = self.compute_strength(stresses)
        row_cells[1].text = str(int(np.round(avg_strength)))
        row_cells[2].text = str(int(np.round(std_strength)))
        # Stress at break
        row_cells = table.add_row().cells
        row_cells[0].text = 'Flexural stress at break (MPa)'
        stresses, strains = self.compute_stress_strain()
        avg_break, std_break = self.compute_stress_at_break(stresses)
        row_cells[1].text = str(int(np.round(avg_break)))
        row_cells[2].text = str(int(np.round(std_break)))

        # Sample specific properties
        p = document.add_paragraph()
        p.add_run('Samples properties test results').bold = True
        table = document.add_table(rows=1, cols=3, style='Table Grid')
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Sample number'
        header_cells[1].text = 'Rupture below 5% strain'
        header_cells[2].text = 'Rupture behaviour'
        behaviour_5pr100 = self.compute_strain_at_break_5p100(strains)
        for i, t in enumerate(self.tests):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i+1)
            row_cells[1].text = behaviour_5pr100[i]
            if i  == 0:
                row_cells[2].text = 'Describe rupture behaviour here'
            else:
                row_cells[2].text = ''

        # Document save
        filename = 'report_examples/ASTM_D790_test_report.docx'
        i = 1
        while filename in os.listdir():
            filename = f'ASTM_D790_test_report_{i}.docx'
            i += 1
        document.save(filename)
        self.report_file = filename


class D3039(test_report):
    """ Class for computing the results of the ASTMD3039 test standard
    D30 Committee. (2017). Test Method for Tensile Properties of Polymer Matrix Composite Materials.
    ASTM International. https://doi.org/10.1520/D3039_D3039M-17
    """

    reference = "D30 Committee. (2017). Test Method for Tensile Properties of Polymer Matrix Composite Materials. " \
                "ASTM International. https://doi.org/10.1520/D3039_D3039M-17 "

    def __init__(self, filenames, widths, thicknesses, lengths, mtr_name="", extensiometer_length=50.8):
        """
        Constructor class for the ASTM D3039 test standard
        :param filenames: filenames for the test data in a list
        :param widths: List of the sample widths in millimeters
        :param thicknesses: List of the sample thicknesses in millimeters
        :param lengths: List of the sample lengths in millimeteres
        :param mtr_name: String, Name of the material
        :param extensiometer_length: float, length of the exensiometer used during the tests
        This class does all the calculations needed for an ASTM D790 test report and writes the results to file
        """
        # Store the attributes
        self.name = mtr_name
        self.widths = widths  # widths of the samples (list)
        self.thickness = thicknesses  # depths of the samples (list)
        self.lengths = lengths  # span used to test the samples (float)
        self.ext_len = extensiometer_length
        self.modulus_upper_strain_range = 0.003
        # Read the test data from the files
        self.tests = self.get_test_data(filenames)
        self.write_report()

    def compute_stress_strain(self):
        """
        Method computing the stress-strain data from the tests results
        Stress formula is from Eq. 6 in ASTM D3039 section 13.1
        Strain formula is from Eq. 7 in ASTM D3039 section 13.2
        :return: stress and strain for all the tests
        """
        all_stress = []
        all_strain = []
        for test, width, thickness in zip(self.tests, self.widths, self.thickness):
            stress = np.array(test['Load']/(width*thickness))
            strain = np.array(test["Extensometer"]/self.ext_len)
            strain = strain[:np.argmax(stress)]
            stress = stress[:np.argmax(stress)]
            all_stress.append(stress)
            all_strain.append(strain)
        return all_stress, all_strain

    def compute_chord_modulus(self, all_values=False):
        """
        Compute the chord modulus from Eq. 8 Section 13.3.1 in standard ASTM D3039
        :return:
        """
        all_stress, all_strain = self.compute_stress_strain()
        std_brk_strain, _ = self.compute_break_strain(all_strain)
        if std_brk_strain < 0.006:
            # "For materials that fail below 6000 μe, a strain range of 25 to 50 % of ultimate is recommended."
            # Table 3, p.9 ASTM D3039
            self.modulus_upper_strain_range = std_brk_strain*0.375

        modulus = []
        #  Compute the modulus for each test
        for stress, strain in zip(all_stress, all_strain):
            # find closest indexes to 0.001 and 0.003 in strains
            # As per table 3 in ASTM D3039
            idx_1 = np.arange(strain.shape[0])[strain >= 0.001][0]
            idx_2 = np.arange(strain.shape[0])[strain >= self.modulus_upper_strain_range][0]

            # Eq. 8 in ASTM D3039 Section 13
            E_chord = (stress[idx_2] - stress[idx_1])/(strain[idx_2] - strain[idx_1])
            modulus.append(E_chord)
        modulus = np.array(modulus)

        if all_values:
            return modulus
        else:
            return modulus.mean(), modulus.std()

    def write_report(self):
        """
        Method writing the ASTM D3039 test report to a .docx (microsoft word) file
        :return: None
        """
        # Compute the stress strain values
        all_stress, all_strain = self.compute_stress_strain()

        # Document header
        document = Document()
        document.add_heading('ASTM D3039 Test report', 1)
        document.add_heading('Montreal ' + date.today().isoformat(), 3)
        document.add_paragraph(self.reference)
        document.add_paragraph('Test done by: Your name')
        document.add_paragraph('Comment on anomalies or problems encountered during testing: Write here')

        # Material identification
        document.add_heading('Material Identification', 2)
        document.add_paragraph(f'Material name: {self.name}', style='List Bullet')
        document.add_paragraph('Material type: ', style='List Bullet')
        document.add_paragraph('Material source: ', style='List Bullet')
        document.add_paragraph('Material code number: ', style='List Bullet')
        document.add_paragraph('Date of material certification: ', style='List Bullet')
        document.add_paragraph('Date of expiration for material certification: ', style='List Bullet')
        document.add_paragraph('Filament diameter: ', style='List Bullet')
        document.add_paragraph('Tow or yarn filament count and twist: ', style='List Bullet')
        document.add_paragraph('Form or weave: ', style='List Bullet')
        document.add_paragraph('Fibre areal weight: ', style='List Bullet')
        document.add_paragraph('Matrix type: ', style='List Bullet')
        document.add_paragraph('Prepreg matrix content: ', style='List Bullet')
        document.add_paragraph('Prepreg volatile content: ', style='List Bullet')
        document.add_paragraph('Composite material ply stacking sequence: ', style='List Bullet')

        # Specimen preparation
        document.add_heading('Specimen preparation', 2)
        document.add_paragraph('Fabrication start date: ', style='List Bullet')
        document.add_paragraph('Fabrication end date: ', style='List Bullet')
        document.add_paragraph('Process specification: ', style='List Bullet')
        document.add_paragraph('Cure cycle: ', style='List Bullet')
        document.add_paragraph('Consolidation method: ', style='List Bullet')
        document.add_paragraph('Description of the equipement used: ', style='List Bullet')
        document.add_paragraph('Average ply thickness (mm): ', style='List Bullet')
        document.add_paragraph('Specimen coupon cutting method: ', style='List Bullet')
        document.add_paragraph('identification of tab geometry, tab material, and tab adhesive used: ',
                               style='List Bullet')
        document.add_paragraph('Direction of cutting and loading of the specimens: ', style='List Bullet')
        document.add_paragraph('Conditioning procedure: ', style='List Bullet')
        document.add_page_break()

        # Specimen dimensions
        p = document.add_paragraph()
        p.add_run('Dimensions of each test specimen').bold = True
        table = document.add_table(rows=1, cols=4, style='Table Grid')
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Specimen number'
        header_cells[1].text = 'Width (mm)'
        header_cells[2].text = 'Thickness (mm)'
        header_cells[3].text = 'Length (mm)'
        for id, b, t, l in zip(range(1, len(self.widths)+1), self.widths, self.thickness, self.lengths):
            row_cells = table.add_row().cells
            row_cells[0].text = str(id)
            row_cells[1].text = str(b)
            row_cells[2].text = str(t)
            row_cells[3].text = str(l)

        # Testing procedure
        document.add_heading('Testing procedure', 2)
        document.add_paragraph('Type of test machine, grips, jaws, grip pressure: ', style='List Bullet')
        document.add_paragraph('Relative humidity and temperature of the testing laboratory: ', style='List Bullet')
        document.add_paragraph('Environment of the test machine environmental chamber (if used): ',
                               style='List Bullet')
        document.add_paragraph(f'Number of specimens tested: {len(self.tests)}', style='List Bullet')
        rate = self.compute_cross_head_rate(self.tests)
        document.add_paragraph(f'Rate of loading head displacement: {rate} mm/min', style='List Bullet')

        # Stress-strain plot section
        document.add_heading('Flexural stress-strain curves', 2)
        self.make_stress_strain_plot('stress-strain_temp', all_strain=all_strain, all_stress=all_stress)
        document.add_picture('stress-strain_temp.png', width=Inches(6))
        document.add_page_break()
        os.remove('stress-strain_temp.png')

        # Test results section
        document.add_heading('Test results', 2)
        document.add_paragraph('Was a speciment rejected after testing: Yes, No, why ?', style='List Bullet')
        p = document.add_paragraph()
        p.add_run('Average properties test results').bold = True
        table = document.add_table(rows=1, cols=3, style='Table Grid')
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Property'
        header_cells[1].text = 'Mean value'
        header_cells[2].text = 'CV %'

        # Modulus
        row_cells = table.add_row().cells
        row_cells[0].text = 'Tangent bending modulus (MPa)'
        avg_modulus, std_modulus = self.compute_chord_modulus()
        row_cells[1].text = str(int(np.round(avg_modulus)))
        row_cells[2].text = str(int(np.round(100*std_modulus/avg_modulus)))

        # Strength
        row_cells = table.add_row().cells
        row_cells[0].text = 'Flexural strength (MPa)'
        avg_strength, std_strength = self.compute_strength(all_stress)
        row_cells[1].text = str(int(np.round(avg_strength)))
        row_cells[2].text = str(int(np.round(100*std_strength/avg_strength)))

        # Failure strain
        row_cells = table.add_row().cells
        row_cells[0].text = 'Strain at break (MPa)'
        avg_brk_strain, std_brk_strain = self.compute_strength(all_strain)
        row_cells[1].text = str(np.around(avg_brk_strain, 3))
        row_cells[2].text = str(int(np.round(100*std_brk_strain/avg_brk_strain)))

        document.add_paragraph(f'Strain range used for chord modulus determination : 0.0'
                               f'01, {self.modulus_upper_strain_range}', style='List Bullet')

        # Sample specific properties
        p = document.add_paragraph()
        p.add_run('Samples properties test results').bold = True
        table = document.add_table(rows=1, cols=5, style='Table Grid')
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Sample number'
        header_cells[1].text = 'Strength (MPa)'
        header_cells[2].text = 'Failure strain (mm/mm)'
        header_cells[3].text = 'Modulus (MPa)'
        header_cells[4].text = 'Rupture behaviour'
        modulus_values = self.compute_chord_modulus(all_values=True)
        stress_values = self.compute_strength(all_stress, all_values=True)
        strain_values = self.compute_break_strain(all_strain, all_values=True)
        for i, sig, eps, mod in zip(range(len(self.tests)), stress_values, strain_values, modulus_values):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i+1)
            row_cells[1].text = str(int(np.around(sig)))
            row_cells[2].text = str(np.around(eps, 3))
            row_cells[3].text = str(int(np.round(mod)))
            if i == 0:
                row_cells[4].text = 'Ruptutre mode codes are shown on figure 4 ASTM D3039'
            else:
                row_cells[4].text = ''

        # Document save
        filename = 'report_examples/ASTM_D3039_test_report.docx'
        i = 1
        while filename in os.listdir():
            filename = f'ASTM_D3039_test_report_{i}.docx'
            i += 1
        document.save(filename)
        self.report_file = filename


class D5868(test_report):
    """ Class for computing the results of the ASTMD5868 test standard
    D14 Committee. (2001). Test Method for Lap Shear Adhesion for Fiber Reinforced Plastic (FRP) Bonding.
    ASTM International. https://doi.org/10.1520/D5868-01R14
    """

    reference = """ D14 Committee. (2001). Test Method for Lap Shear Adhesion for Fiber Reinforced Plastic (FRP) 
    Bonding. ASTM International. https://doi.org/10.1520/D5868-01R14"""

    def __init__(self, filenames, areas, mtr_name=""):
        """ Constructor class for the ASTM D5868 test report"""
        # Store attributes
        self.name = mtr_name  # Material name
        self.areas = areas
        # load data and write report
        self.tests = self.get_test_data(filenames)
        self.write_report()

    def compute_stress_times(self):
        """
        Method computing the stress/time data from the tests results
        :return: stress and strain for all the tests
        """
        all_stress, all_times = [], []

        for test, A in zip(self.tests, self.areas):
            stress = np.array(test['Load']/A)
            time = test['Time'][:np.argmax(stress)]
            stress = stress[:np.argmax(stress)]
            all_stress.append(stress)
            all_times.append(time)

        return all_stress, all_times

    def write_report(self):
        """
        Method to write the test report from the test data
        :return: None
        """
        # Compute stress and strain values
        stresses, times = self.compute_stress_times()

        # Document header
        document = Document()
        document.add_heading('ASTM D5868 Test report', 1)
        document.add_heading(date.today().isoformat(), 3)
        document.add_paragraph(self.reference)

        # Adhesive identification
        document.add_heading('Adhesive Identification', 2)
        document.add_paragraph(f'Adhesive name: {self.name}', style='List Bullet')
        document.add_paragraph('Adhesive type: ', style='List Bullet')
        document.add_paragraph('Adhesive source: ', style='List Bullet')
        document.add_paragraph('Adhesive code number: ', style='List Bullet')

        # Substrate identification
        document.add_heading('Substrate Identification', 2)
        document.add_paragraph(f'Substrate name: {self.name}', style='List Bullet')
        document.add_paragraph('Substrate type: ', style='List Bullet')
        document.add_paragraph('Substrate source: ', style='List Bullet')
        document.add_paragraph('Substrate code number: ', style='List Bullet')
        document.add_paragraph('Substrate ply stacking sequence: ', style='List Bullet')

        # Specimen preparation
        document.add_heading('Specimen preparation', 2)
        document.add_paragraph('Adhesive cure time: ', style='List Bullet')
        document.add_paragraph('Adhesive cure temperature: ', style='List Bullet')
        document.add_paragraph('Other adhesive bonding conditions: ', style='List Bullet')
        p = document.add_paragraph()
        p.add_run('Specimens dimensions').bold = True
        table = document.add_table(rows=1, cols=2, style='Table Grid')
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Specimen number'
        header_cells[1].text = 'Area (mm^2)'
        for id, A in zip(range(1, len(self.areas) + 1), self.areas):
            row_cells = table.add_row().cells
            row_cells[0].text = str(id)
            row_cells[1].text = str(np.around(A))

        # Testing procedure
        document.add_heading('Testing procedure', 2)
        document.add_paragraph(f'Test temperature:', style='List Bullet')
        document.add_paragraph(f'Other testing conditions:', style='List Bullet')

        # Stress-strain plot section
        document.add_heading('Flexural stress-strain curves', 2)
        self.make_stress_strain_plot('stress-strain_temp', all_stress=stresses, all_strain=times)
        document.add_picture('stress-strain_temp.png', width=Inches(6))
        os.remove('stress-strain_temp.png')

        # Test results section
        document.add_heading('Test results', 2)
        document.add_paragraph('Was a speciment rejected after testing: Yes, No, why ?', style='List Bullet')
        p = document.add_paragraph()
        p.add_run('Average properties test results').bold = True
        table = document.add_table(rows=1, cols=5, style='Table Grid')
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Property'
        header_cells[1].text = 'Mean value'
        header_cells[2].text = 'Standard deviation'
        header_cells[3].text = 'Min value'
        header_cells[4].text = 'Max value'

        # Max load
        row_cells = table.add_row().cells
        row_cells[0].text = 'Peak load values (MPa)'
        avg_stren, std_stren = self.compute_strength(stresses)
        row_cells[1].text = str(np.around(avg_stren, 1))
        row_cells[2].text = str(np.around(std_stren, 1))
        all_stren = self.compute_strength(stresses, all_values=True)
        row_cells[3].text = str(np.around(np.min(all_stren), 1))
        row_cells[4].text = str(np.around(np.max(all_stren), 1))

        # Sample specific properties
        p = document.add_paragraph()
        p.add_run('Samples properties test results').bold = True
        table = document.add_table(rows=1, cols=3, style='Table Grid')
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Sample number'
        header_cells[1].text = 'Peak load (MPa)'
        header_cells[2].text = 'Rupture behaviour'
        for i, sig in enumerate(all_stren):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i + 1)
            row_cells[1].text = str(int(np.around(sig)))
            if i == 0:
                row_cells[2].text = 'Describe rupture behaviour here'
            else:
                row_cells[2].text = ''
        document.add_paragraph(' ')

        # Document save
        filename = 'report_examples/ASTM_D5868_test_report.docx'
        i = 1
        while filename in os.listdir():
            filename = f'ASTM_D5868_test_report_{i}.docx'
            i += 1
        document.save(filename)
        self.report_file = filename
